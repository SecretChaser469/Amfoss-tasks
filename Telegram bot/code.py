import os
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Updater, CommandHandler, CallbackQueryHandler, MessageHandler, Filters
import requests
import csv
from docx import Document

# Bot Token (get it from BotFather on Telegram)
TOKEN = 'YOUR_TELEGRAM_BOT_TOKEN'

# Global reading list storage
reading_list = []

def start(update: Update, context):
    update.message.reply_text("Welcome to PagePal! I'm here to help you find and manage books. Type /help for available commands.")

def help_command(update: Update, context):
    update.message.reply_text("""
    Available commands:
    /start - Start the bot
    /book - Find books by genre
    /preview - Get a preview link for a specific book
    /list - Manage your reading list
    /reading_list - View and manage your reading list
    """)

def book(update: Update, context):
    update.message.reply_text("Please type in the genre of the book you're interested in:")

def handle_genre(update: Update, context):
    genre = update.message.text
    book_data = get_books_by_genre(genre)
    
    # Generate CSV file with book details
    filename = f'{genre}_books.csv'
    with open(filename, 'w', newline='') as csvfile:
        fieldnames = ['Title', 'Author', 'Description', 'Year', 'Language', 'Preview Link']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for book in book_data:
            writer.writerow(book)

    update.message.reply_document(open(filename, 'rb'))
    os.remove(filename)

def get_books_by_genre(genre):
    # Example mock data; you should connect to an API like Google Books or OpenLibrary
    books = [
        {"Title": "Book 1", "Author": "Author 1", "Description": "Description 1", "Year": 1999, "Language": "English", "Preview Link": "http://example.com/preview1"},
        {"Title": "Book 2", "Author": "Author 2", "Description": "Description 2", "Year": 2001, "Language": "English", "Preview Link": "http://example.com/preview2"},
    ]
    return books

def preview(update: Update, context):
    update.message.reply_text("Please type in the name of the book you want a preview for:")

def handle_preview(update: Update, context):
    book_name = update.message.text
    preview_link = get_preview_link(book_name)
    update.message.reply_text(f'Preview link for {book_name}: {preview_link}')

def get_preview_link(book_name):
    # Example mock data; fetch real data from an API
    return "http://example.com/preview"

def reading_list(update: Update, context):
    keyboard = [
        [InlineKeyboardButton("Add a book", callback_data='add')],
        [InlineKeyboardButton("Delete a book", callback_data='delete')],
        [InlineKeyboardButton("View Reading List", callback_data='view')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    update.message.reply_text('What would you like to do?', reply_markup=reply_markup)

def handle_reading_list(update: Update, context):
    query = update.callback_query
    query.answer()

    if query.data == 'add':
        query.edit_message_text(text="Please type in the book title to add:")
        context.user_data['action'] = 'add'
    elif query.data == 'delete':
        query.edit_message_text(text="Please type in the book title to delete:")
        context.user_data['action'] = 'delete'
    elif query.data == 'view':
        query.edit_message_text(text="Generating your reading list...")
        generate_reading_list(query.message.chat_id)

def add_book_to_list(book_title):
    preview_link = get_preview_link(book_title)
    reading_list.append({"Title": book_title, "Preview Link": preview_link})

def delete_book_from_list(book_title):
    global reading_list
    reading_list = [book for book in reading_list if book['Title'] != book_title]

def generate_reading_list(chat_id):
    doc = Document()
    doc.add_heading('Reading List', 0)

    for book in reading_list:
        doc.add_paragraph(f"Title: {book['Title']}")
        doc.add_paragraph(f"Preview Link: {book['Preview Link']}")
        doc.add_paragraph('\n')

    filename = f'reading_list_{chat_id}.docx'
    doc.save(filename)

    context.bot.send_document(chat_id=chat_id, document=open(filename, 'rb'))
    os.remove(filename)

def handle_message(update: Update, context):
    user_action = context.user_data.get('action', None)
    
    if user_action == 'add':
        book_title = update.message.text
        add_book_to_list(book_title)
        update.message.reply_text(f'{book_title} has been added to your reading list.')
        context.user_data['action'] = None
    elif user_action == 'delete':
        book_title = update.message.text
        delete_book_from_list(book_title)
        update.message.reply_text(f'{book_title} has been removed from your reading list.')
        context.user_data['action'] = None

def main():
    updater = Updater(TOKEN, use_context=True)
    dp = updater.dispatcher

    dp.add_handler(CommandHandler("start", start))
    dp.add_handler(CommandHandler("help", help_command))
    dp.add_handler(CommandHandler("book", book))
    dp.add_handler(CommandHandler("preview", preview))
    dp.add_handler(CommandHandler("reading_list", reading_list))
    dp.add_handler(MessageHandler(Filters.text & ~Filters.command, handle_message))
    dp.add_handler(CallbackQueryHandler(handle_reading_list))

    updater.start_polling()
    updater.idle()

if __name__ == '__main__':
    main()
